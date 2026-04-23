from docx import Document
from fpdf import FPDF
import os

os.makedirs('samples', exist_ok=True)

# Create a sample DOCX resume
resume = Document()
resume.add_heading('Alicia Morgan', level=0)
resume.add_paragraph('Professional Summary')
resume.add_paragraph('Experienced data analyst with a proven record of improving reporting processes and delivering insights for cross-functional teams.')
resume.add_paragraph('Skills')
resume.add_paragraph('Python, SQL, Tableau, Power BI, machine learning, data visualization, project management')
resume.add_paragraph('Experience')
resume.add_paragraph('Senior Data Analyst at TechWave, 2021 - Present')
resume.add_paragraph('Improved customer reporting accuracy by 22% and optimized dashboard performance for stakeholders.')
resume.add_paragraph('Education')
resume.add_paragraph('B.Sc. in Information Systems, State University')
resume.add_paragraph('Projects')
resume.add_paragraph('Developed a predictive model for churn reduction, increasing retention by 15%.')
resume.save('samples/sample_resume.docx')

# Create a sample PDF resume
pdf = FPDF()
pdf.set_auto_page_break(auto=True, margin=15)
pdf.add_page()
pdf.set_font('Arial', 'B', 16)
pdf.cell(0, 10, 'Alicia Morgan', ln=True)
pdf.ln(2)
pdf.set_font('Arial', '', 12)
lines = [
    'Professional Summary',
    'Experienced data analyst with a proven record of improving reporting processes and delivering insights for cross-functional teams.',
    '',
    'Skills',
    'Python, SQL, Tableau, Power BI, machine learning, data visualization, project management',
    '',
    'Experience',
    'Senior Data Analyst at TechWave, 2021 - Present',
    'Improved customer reporting accuracy by 22% and optimized dashboard performance for stakeholders.',
    '',
    'Education',
    'B.Sc. in Information Systems, State University',
    '',
    'Projects',
    'Developed a predictive model for churn reduction, increasing retention by 15%.',
]
for line in lines:
    pdf.multi_cell(0, 8, line)
pdf.output('samples/sample_resume.pdf')
print('Sample resumes created in samples/')
